## Module --> pip install python-docx

import docx

dcmt = doc.Document('example.docx')

dcmt.paragraphs     #Returns a list of paragraphs objects

dcmt.paragraphs[0].text     #Returns text from de paragraph object

p = dcmt.paragraphs[1]

p.runs          # returns a list

p.runs[1].text  # returns text

p.runs[1].bold  # returns text_format == 'bold' --> posibilities (bold, italic, underline)

p.runs[1].underline = True  # underlines text

p.runs[1].text = 'Changing text'

dcmt.save('example.docx')   #Commiting changes

p.style = 'Title'


d = docx.Document()     #Creates a docx document in memory(not hard drive).

d.add_paragraph('Hello, this is a paragraph.')
d.add_paragraph('This is another paragraph.')

d.save('example2.docx')         # <-- Insert path to save.

p = d.paragraphs[0]

p.add_run('This is a new run.')

p.runs[1].bold = True       # #1 run now is bold.

d.save('example2.docx')


#=== Copying a document function ===#

def getText(filename):
    '''pre-condition: word document\'s path
    returns: string wich contains text from the document.
    '''
    doc = docx.Document(filename)
    fullText = []

    for para in doc.paragraphs:
        fullText.append(para.text)
    
    return '\n.'.join(fullText)

